from __future__ import annotations

import io
import logging
import re

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def _extract_pdf_text(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    return "\n".join(text_parts)


def _extract_txt_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def _clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    try:
        lower_name = filename.lower()
        if lower_name.endswith(".pdf"):
            raw_text = _extract_pdf_text(file_bytes)
        elif lower_name.endswith(".docx"):
            raw_text = _extract_docx_text(file_bytes)
        elif lower_name.endswith(".txt"):
            raw_text = _extract_txt_text(file_bytes)
        else:
            logger.warning("Unsupported file type for %s", filename)
            return ""
        return _clean_extracted_text(raw_text)
    except Exception as exc:
        logger.warning("Failed to extract text from %s: %s", filename, exc)
        return ""



# def extract_text_new(file_bytes : bytes , filename : str ) -> str :
#     ans = []
#     for i in range(0 , 10):
#         ans.append("text_recieved")
#     return ""
